import os
import io
import json
import logging
from typing import Optional
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from groq import Groq

from complaint_shared import FIELD_KEYS
from agent_graph import build_agent_graph, run_extraction

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("ai-voa")


app = FastAPI(title="AI-VOA Complaint System Backend", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
if not GROQ_API_KEY:
    logger.warning(
        "GROQ_API_KEY is not set. Set it as an environment variable "
        "(get one from https://console.groq.com/keys). Extraction calls will fail until then."
    )

client = Groq(api_key=GROQ_API_KEY)
MODEL_NAME = os.environ.get("GROQ_MODEL", "openai/gpt-oss-20b")

# Compiled once at import time; reused (cheaply) across requests.
AGENT_GRAPH = build_agent_graph(client, MODEL_NAME)

MAX_UPLOAD_SIZE = 10 * 1024 * 1024


class ComplaintInput(BaseModel):
    text: str
    mode: Optional[str] = None
    current_data: Optional[dict] = None


class ComplaintRecord(BaseModel):
    complaintSource: str = ""
    customerName: str = ""
    productName: str = ""
    productStrength: str = ""
    batchNumber: str = ""
    manufacturingDate: str = ""
    expiryDate: str = ""
    quantityAffected: str = ""
    complaintType: str = ""
    complaintDate: str = ""
    description: str = ""
    initialSeverity: str = "Medium"
    priority: str = "Medium"
    complaintCategory: str = ""
    suggestedSeverity: str = ""
    suggestedNextAction: str = ""
    initialRiskAssessment: str = ""
    status: str = "Pending Triage"


DB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "complaints_db.json")


def _load_db():
    if os.path.exists(DB_FILE):
        try:
            with open(DB_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
                return data.get("complaints", []), data.get("next_id", 1)
        except Exception as e:
            logger.warning("Could not read %s (%s); starting with an empty store", DB_FILE, e)
    return [], 1


def _save_db():
    try:
        with open(DB_FILE, "w", encoding="utf-8") as f:
            json.dump({"complaints": COMPLAINTS_DB, "next_id": _next_id}, f, indent=2)
    except Exception as e:
        logger.error("Failed to persist complaints_db.json: %s", e)


COMPLAINTS_DB, _next_id = _load_db()


@app.get("/api/health")
def read_root():
    return {
        "message": "AI-VOA Complaint System Backend Running.",
        "model": MODEL_NAME,
        "agent_framework": "langgraph",
        "groq_key_configured": bool(GROQ_API_KEY),
    }


def _extract_pdf_text(contents: bytes) -> str:
    """Requires: pip install pypdf"""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(contents))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(pages).strip()


def _extract_docx_text(contents: bytes) -> str:
    """Requires: pip install python-docx"""
    from docx import Document

    doc = Document(io.BytesIO(contents))
    chunks = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return "\n".join(chunks).strip()


@app.post("/api/extract-complaint")
async def extract_complaint(data: ComplaintInput):
    if not data.text or not data.text.strip():
        raise HTTPException(status_code=400, detail="No complaint text provided.")

    if not GROQ_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="Server misconfigured: GROQ_API_KEY environment variable is not set.",
        )

    try:
        return run_extraction(
            AGENT_GRAPH,
            text=data.text,
            mode=data.mode,
            current_data=data.current_data,
        )
    except json.JSONDecodeError as e:
        logger.error("Model did not return valid JSON: %s", e)
        raise HTTPException(
            status_code=502,
            detail="AI model returned an unparseable response. Try again or shorten the input text.",
        )
    except Exception as e:
        logger.error("Extraction Error: %s", e)
        raise HTTPException(status_code=502, detail=f"AI extraction failed: {e}")


@app.post("/api/extract-complaint-file")
async def extract_complaint_file(file: UploadFile = File(...)):
    """Handles PDF/DOCX uploads from the sidebar's drag-and-drop zone.
    (TXT/EML are read as plain text in the browser and go through
    /api/extract-complaint instead — they never hit this endpoint.)
    """
    if not GROQ_API_KEY:
        raise HTTPException(
            status_code=500,
            detail="Server misconfigured: GROQ_API_KEY environment variable is not set.",
        )

    filename = file.filename or "uploaded file"
    ext = os.path.splitext(filename)[1].lower()

    if ext not in (".pdf", ".docx"):
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. This endpoint handles PDF and DOCX only.",
        )

    contents = await file.read()

    if len(contents) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=400, detail=f"'{filename}' is over the 10MB limit.")

    try:
        text = _extract_pdf_text(contents) if ext == ".pdf" else _extract_docx_text(contents)
    except ImportError as e:
        logger.error("Missing text-extraction dependency: %s", e)
        raise HTTPException(
            status_code=500,
            detail=(
                "Server is missing a required package for reading this file type. "
                "Install with: pip install pypdf python-docx"
            ),
        )
    except Exception as e:
        logger.error("Failed to parse '%s': %s", filename, e)
        raise HTTPException(status_code=400, detail=f"Could not read '{filename}': {e}")

    if not text.strip():
        raise HTTPException(
            status_code=400,
            detail=(
                f"No extractable text found in '{filename}'. It may be a scanned/"
                "image-only document — try a text-based PDF or DOCX instead."
            ),
        )

    try:
        return run_extraction(AGENT_GRAPH, text=text, mode=None, current_data=None)
    except json.JSONDecodeError as e:
        logger.error("Model did not return valid JSON: %s", e)
        raise HTTPException(status_code=502, detail="AI model returned an unparseable response. Try again.")
    except Exception as e:
        logger.error("Extraction Error: %s", e)
        raise HTTPException(status_code=502, detail=f"AI extraction failed: {e}")


@app.post("/api/complaints")
async def save_complaint(record: ComplaintRecord):
    global _next_id
    complaint = record.dict()
    complaint["id"] = _next_id
    complaint["status"] = complaint.get("status") or "Pending Triage"
    COMPLAINTS_DB.append(complaint)
    _next_id += 1
    _save_db()
    logger.info("Saved complaint #%s", complaint["id"])
    return complaint


@app.get("/api/complaints")
async def list_complaints():
    return COMPLAINTS_DB


@app.get("/api/complaints/{complaint_id}")
async def get_complaint(complaint_id: int):
    for complaint in COMPLAINTS_DB:
        if complaint["id"] == complaint_id:
            return complaint
    raise HTTPException(status_code=404, detail=f"Complaint #{complaint_id} not found.")


@app.post("/api/complaints/{complaint_id}/commit")
async def commit_complaint(complaint_id: int):
    """Called when the user drags the 'Commit to QMS Ledger' slider to
    confirm. Marks the complaint as officially logged."""
    for complaint in COMPLAINTS_DB:
        if complaint["id"] == complaint_id:
            complaint["status"] = "Committed to QMS Ledger"
            _save_db()
            logger.info("Committed complaint #%s to QMS ledger", complaint_id)
            return complaint
    raise HTTPException(status_code=404, detail=f"Complaint #{complaint_id} not found.")


# --- Serve React Frontend ---
frontend_dist = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend", "dist")

if os.path.exists(frontend_dist):
    app.mount("/assets", StaticFiles(directory=os.path.join(frontend_dist, "assets")), name="assets")

    @app.get("/{full_path:path}")
    async def serve_react(full_path: str):
        if full_path.startswith("api/") or full_path in ("docs", "redoc", "openapi.json"):
            raise HTTPException(status_code=404, detail="Not found")
        
        index_file = os.path.join(frontend_dist, "index.html")
        if os.path.exists(index_file):
            return FileResponse(index_file)
        return {"error": "Frontend build not found"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("server:app", host="0.0.0.0", port=8000, reload=True)